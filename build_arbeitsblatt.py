"""
Generiert das Arbeitsblatt zur TikTok-AGB-Praesentation.
Zwei Versionen: Schueler-AB (leer) und Lehrer-Loesungsblatt.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Farben (fuer Papier lesbar abgedunkelt)
TT_CYAN = RGBColor(0x0C, 0x8C, 0x88)   # Teal statt grellem Cyan
TT_RED = RGBColor(0xE1, 0x1D, 0x48)    # TikTok-Rot
TT_PURPLE = RGBColor(0x7C, 0x3A, 0xED)
TT_GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK = RGBColor(0x1a, 0x1a, 0x1a)
GREY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x16, 0x80, 0x3C)


def set_margins(doc, top=1.5, bottom=1.5, left=1.8, right=1.8):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_para(doc, text, size=11, bold=False, italic=False, color=DARK, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_lines(doc, count=2, spacing=2.0, keep_together=False):
    """Schreiblinien hinzufuegen - leere Zeilen mit hellem Unterstrich.
    spacing = Zeilenabstand (2.0 = doppelt, damit genug Platz zum Schreiben bleibt).
    keep_together = bindet die Linien aneinander (kein Seitenumbruch mittendrin)."""
    for i in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = spacing
        if keep_together and i < count - 1:
            p.paragraph_format.keep_with_next = True
        run = p.add_run("_" * 90)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def set_row_cant_split(row):
    """Verhindert, dass eine Tabellenzeile ueber einen Seitenumbruch zerrissen wird."""
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    trPr.append(cant)


def add_field(paragraph, field_code, size=8, color=GREY):
    """Fuegt ein Word-Feld (z.B. PAGE / NUMPAGES) als rohes XML in einen Absatz ein.
    Word berechnet den Wert beim Oeffnen/Drucken selbst."""
    run = paragraph.add_run()
    run.font.size = Pt(size)
    run.font.color.rgb = color
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field_code
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_page_number_footer(doc):
    """Seitenzahl im Format 'Seite X von Y' zentriert in die Fusszeile (auf jeder Seite)."""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Seite ")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    add_field(p, "PAGE")
    r2 = p.add_run(" von ")
    r2.font.size = Pt(8)
    r2.font.color.rgb = GREY
    add_field(p, "NUMPAGES")


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_kopfzeile(doc, titel, untertitel):
    """Logo-freie Kopfzeile mit Titel + Name/Klasse/Datum-Zeilen."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(titel)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = DARK
    r.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(untertitel)
    r2.font.size = Pt(11)
    r2.font.italic = True
    r2.font.color.rgb = GREY

    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    widths = [Cm(7), Cm(4), Cm(5)]
    for i, w in enumerate(widths):
        t.columns[i].width = w
    cells = t.rows[0].cells
    for i, label in enumerate(["Name: _____________________________",
                                "Klasse: __________",
                                "Datum: ___________"]):
        cells[i].width = widths[i]
        para = cells[i].paragraphs[0]
        run = para.add_run(label)
        run.font.size = Pt(10)
        run.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def aufgabe_header(doc, nr, titel, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True  # Titel klebt am Aufgabentext (kein verwaister Header)
    r1 = p.add_run(f"Aufgabe {nr}  -  ")
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = color
    r1.font.name = "Calibri"
    r2 = p.add_run(titel)
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.color.rgb = DARK
    r2.font.name = "Calibri"


def baue_arbeitsblatt(loesungen=False):
    doc = Document()
    set_margins(doc, top=1.2, bottom=1.2)

    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if loesungen:
        add_kopfzeile(doc,
            "Lösungsblatt: Was du wirklich zustimmst (TikTok)",
            "TikTok-AGB - Lehrer-Version mit Erwartungshorizont")
    else:
        add_kopfzeile(doc,
            "Was du wirklich zustimmst - TikTok",
            "TikTok-AGB - Begleitheft zur Präsentation")

    # ──────────────────────────────────────────────────────────
    # AUFGABE 1 - Schaetzen vor der Praesentation
    # ──────────────────────────────────────────────────────────
    aufgabe_header(doc, 1, "Schätzen - bevor die Präsentation startet", color=TT_CYAN)
    add_para(doc, "Wie schätzt du es ein? Kreuze an oder schreibe deine Vermutung. Wir lösen es danach gemeinsam auf.",
             italic=True, color=GREY, space_after=10)

    items = [
        ("Ab welchem Alter darfst du laut TikTok-AGB live streamen und Geschenke (Geld) bekommen?",
         "[ ] ab 13   [ ] ab 16   [ ] ab 18   [ ] egal",
         "Ab 18 (TikTok LIVE). Zur Nutzung allgemein ab 13, Direktnachrichten ab 16."),
        ("Wenn dir ein Fan bei LIVE ein Geschenk schickt - wie viel vom Wert bekommst du als Creator?",
         "[ ] fast alles   [ ] etwa die Hälfte   [ ] nur ein Viertel",
         "Etwa die Hälfte - TikTok behält rund 50 %."),
        ("Was sammelt TikTok automatisch über dich? (Mehrfachnennung möglich)",
         "[ ] Standort   [ ] Tipprhythmus   [ ] In-App-Browsing   [ ] nichts ohne Erlaubnis",
         "Standort, Tipprhythmus und In-App-Browsing - alles automatisch."),
        ("Wie hoch war 2025 die EU-Strafe gegen TikTok (Daten nach China)?",
         "[ ] 5 Mio. EUR   [ ] 53 Mio. EUR   [ ] 530 Mio. EUR   [ ] 5 Mrd. EUR",
         "530 Mio. EUR - verhängt von der irischen Datenschutzbehörde am 2. Mai 2025."),
        ("Verdient TikTok an deinen Videos Geld, ohne dich zu bezahlen?",
         "[ ] Ja   [ ] Nein   [ ] nur mit meiner Erlaubnis",
         "Ja. Die Lizenz ist gebührenfrei, und deine Inhalte trainieren TikToks KI."),
    ]
    for q, opts, ans in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"  >  {q}")
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = DARK
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(opts)
        r2.font.size = Pt(11)
        r2.font.color.rgb = GREY
        if loesungen:
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Cm(0.5)
            p3.paragraph_format.space_after = Pt(6)
            r3 = p3.add_run("Lösung:  " + ans)
            r3.font.size = Pt(10)
            r3.font.bold = True
            r3.font.color.rgb = GREEN

    # ──────────────────────────────────────────────────────────
    # AUFGABE 2 - Notizen waehrend der Praesentation (aktiv, eigene Worte)
    # ──────────────────────────────────────────────────────────
    aufgabe_header(doc, 2, "Mitdenken während der Präsentation", color=TT_RED)
    add_para(doc, "Schreibe pro Bereich in EIGENEN Worten mindestens 2 Dinge auf, die dich überraschen - "
                  "nicht abschreiben, sondern so, dass deine Banknachbarin / dein Banknachbar es versteht.",
             italic=True, color=GREY, space_after=8)

    bereiche = [
        ("Deine Videos (Lizenz & KI):", TT_CYAN, [
            "Lizenz: weltweit, gebührenfrei, übertragbar, unterlizenzierbar - du bekommst nichts.",
            "Deine Inhalte (Gesicht, Stimme, gesprochene Wörter) trainieren TikToks KI/Algorithmen.",
        ]),
        ("Algorithmus & Daten:", TT_RED, [
            "Automatisch gesammelt: Tipprhythmus, Standort, In-App-Browsing, Verweildauer.",
            "Der 'Für dich'-Feed ist endlos und so gebaut, dass man möglichst lange dranbleibt.",
        ]),
        ("Datenweitergabe:", TT_PURPLE, [
            "Konzern-Firmen im Ausland und Behörden können Zugriff bekommen.",
            "2025: 530 Mio. EUR Strafe, weil EU-Daten aus China abrufbar/dort gespeichert waren.",
        ]),
        ("Geld:", TT_GOLD, [
            "LIVE-Geschenke = echtes Geld über Coins; TikTok behält rund die Hälfte.",
            "Account löschen heißt nicht 'alles weg' - die Lizenz bleibt bestehen.",
        ]),
    ]
    for titel, farbe, punkte in bereiche:
        bh = add_para(doc, titel, bold=True, color=farbe, size=12, space_after=4, space_before=6)
        if loesungen:
            for line in punkte:
                add_para(doc, "- " + line, size=10, color=DARK, space_after=2)
        else:
            bh.paragraph_format.keep_with_next = True  # Bereich-Titel + Linien zusammenhalten
            add_lines(doc, 2, keep_together=True)

    # ──────────────────────────────────────────────────────────
    # AUFGABE 3 - Tabelle Mythos vs. AGB-Realitaet (Partnerarbeit)
    # ──────────────────────────────────────────────────────────
    aufgabe_header(doc, 3, "Mythos vs. AGB-Realität  (Partnerarbeit)", color=TT_CYAN)
    add_para(doc, "Think-Pair-Share: Fülle die rechte Spalte zuerst ALLEIN aus. Vergleicht dann zu zweit "
                  "und einigt euch auf die beste Formulierung. Links steht, was viele glauben.",
             italic=True, color=GREY, space_after=8)

    headers = ["Das denken viele", "Was die AGB / Realität sagen"]
    rows_data = [
        ("Meine Videos gehören nur mir.",
         "Ja - aber TikTok bekommt eine kostenlose, weltweite, übertragbare Lizenz darauf."),
        ("TikTok verdient nichts an meinen Videos.",
         "Doch: Lizenz gebührenfrei für dich; deine Inhalte trainieren die KI."),
        ("TikTok weiß nur, was ich like.",
         "Auch Tipprhythmus, Standort, In-App-Browsing und jede Sekunde Verweildauer."),
        ("Meine Daten bleiben in Europa.",
         "Zugriff aus China möglich - 530 Mio. EUR EU-Strafe 2025."),
        ("LIVE-Geschenke sind nur Spaß.",
         "Echtes Geld; TikTok behält rund 50 %; empfangen erst ab 18."),
        ("Account löschen = alles weg.",
         "Nein, nicht sofort/komplett; die Lizenz bleibt bestehen."),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    widths = [Cm(6.5), Cm(10.0)]
    for i, w in enumerate(widths):
        table.columns[i].width = w

    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = widths[i]
        c.text = ""
        para = c.paragraphs[0]
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(c, "E11D48" if i == 0 else "0C8C88")

    for ri, (mythos, real) in enumerate(rows_data, start=1):
        set_row_cant_split(table.rows[ri])
        c0 = table.rows[ri].cells[0]
        c0.text = ""
        c0.width = widths[0]
        p = c0.paragraphs[0]
        r = p.add_run(mythos)
        r.font.size = Pt(10)
        r.font.italic = True

        c1 = table.rows[ri].cells[1]
        c1.text = ""
        c1.width = widths[1]
        p = c1.paragraphs[0]
        if loesungen:
            r = p.add_run(real)
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK
        else:
            # 3 Schreibzeilen Hoehe pro Zelle (genug Platz zum Reinschreiben)
            p.add_run(" ").font.size = Pt(10)
            for _ in range(2):
                ep = c1.add_paragraph()
                ep.paragraph_format.space_after = Pt(0)
                ep.paragraph_format.line_spacing = 1.5
                ep.add_run(" ").font.size = Pt(10)

    # ──────────────────────────────────────────────────────────
    # AUFGABE 4 - AGB-Vokabel-Decoder (Matching) - startet auf neuer Seite
    # ──────────────────────────────────────────────────────────
    if not loesungen:
        doc.add_page_break()
    aufgabe_header(doc, 4, "AGB-Decoder: Fachbegriff trifft Klartext", color=TT_PURPLE)
    add_para(doc, "In der Lizenz (§ 4.9) stehen vier Wörter, die harmlos klingen, aber viel bedeuten. "
                  "Ordne jedem Begriff die richtige Erklärung zu - trage den Buchstaben in das Kästchen ein.",
             italic=True, color=GREY, space_after=8)

    # Begriff -> (Loesungsbuchstabe, Klartext); Erklaerungen werden bewusst gemischt ausgegeben.
    begriffe = [
        ("1.  nicht-exklusiv", "C"),
        ("2.  gebührenfrei", "B"),
        ("3.  übertragbar", "D"),
        ("4.  unterlizenzierbar", "E"),
        ("5.  weltweit", "A"),
    ]
    erklaerungen = [
        ("A", "Gilt überall auf der Erde - nicht nur in Deutschland."),
        ("B", "TikTok zahlt dir dafür keinen Cent."),
        ("C", "Du gibst dein Recht nicht allein ab - du darfst dein Video auch selbst weiternutzen."),
        ("D", "TikTok darf das Recht an andere Firmen weitergeben."),
        ("E", "TikTok darf anderen erlauben, deine Inhalte ebenfalls zu nutzen."),
    ]

    add_para(doc, "Begriffe:", bold=True, size=11, color=DARK, space_after=2, space_before=2)
    for begriff, sol in begriffe:
        kasten = f"[ {sol} ]" if loesungen else "[          ]"
        col = GREEN if loesungen else GREY
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(begriff + "   ")
        r1.font.size = Pt(11)
        r1.font.bold = True
        r1.font.color.rgb = DARK
        r2 = p.add_run(kasten)
        r2.font.size = Pt(11)
        r2.font.bold = True
        r2.font.color.rgb = col

    add_para(doc, "Erklärungen (in falscher Reihenfolge):", bold=True, size=11, color=DARK,
             space_after=2, space_before=6)
    for buchstabe, txt in erklaerungen:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{buchstabe})  ")
        r1.font.size = Pt(10.5)
        r1.font.bold = True
        r1.font.color.rgb = TT_PURPLE
        r2 = p.add_run(txt)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK

    # ──────────────────────────────────────────────────────────
    # AUFGABE 5 - Reflexion
    # ──────────────────────────────────────────────────────────
    aufgabe_header(doc, 5, "Reflexion - deine Meinung zählt", color=TT_GOLD)
    add_para(doc, "Beantworte zwei der folgenden vier Fragen ausführlich (je 3-5 Sätze).",
             italic=True, color=GREY, space_after=8)

    fragen = [
        "1. Würdest du diese AGB unterschreiben, wenn man sie dir auf Papier vorlegt? Begründe.",
        "2. Was machst du ab heute anders auf TikTok? Nenne mindestens zwei konkrete Schritte.",
        "3. Sollte der personalisierte 'Für dich'-Feed für Jugendliche standardmäßig aus sein? Pro/Kontra.",
        "4. Was bedeutet es für dich, dass TikTok deine Videos 'gebührenfrei und weltweit' nutzen darf?",
    ]

    if loesungen:
        erwart = [
            ("1.", "Erwartung: Erkennen, dass niemand einen so langen Vertrag ungelesen unterschreiben würde - Bewusstsein für die Ungleichheit zwischen Nutzer und Konzern."),
            ("2.", "Erwartung: Konkrete Schritte wie Konto auf privat, personalisierten Feed aus, Standort aus, Links extern öffnen, Bildschirmzeit-Limit, keine Coins kaufen, bewusster posten."),
            ("3.", "Erwartung: Pro (Schutz vor Sucht-Schleife, Jugendschutz) und Kontra (Selbstbestimmung, weniger relevante Inhalte). Keine richtige Antwort - bewertet wird die Argumentationstiefe."),
            ("4.", "Erwartung: Verständnis, dass man die kreative Arbeit kostenlos und weltweit an einen Konzern verschenkt, der damit Geld verdienen darf - Wert der eigenen Inhalte."),
        ]
        for nr, txt in erwart:
            add_para(doc, txt, size=10, color=GREY, italic=True, space_after=8)
    else:
        for f in fragen:
            fp = add_para(doc, f, size=11, bold=True, color=DARK, space_after=4)
            fp.paragraph_format.keep_with_next = True  # Frage + ihre Linien zusammenhalten
            add_lines(doc, 3, keep_together=True)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ──────────────────────────────────────────────────────────
    # PROFI-AUFGABE - Rechnen mit dem 50-%-Split (Differenzierung)
    # ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run("Profi-Aufgabe (für Schnelle)  -  ")
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = TT_RED
    r2 = p.add_run("Rechnen mit dem 50-%-Trick")
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.color.rgb = DARK

    add_para(doc, "Bei LIVE-Geschenken behält TikTok rund die Hälfte. Rechne und denke nach:",
             italic=True, color=GREY, space_after=6)

    profi = [
        ("a) Ein Creator bekommt Geschenke im Wert von 80 €. Wie viel bleibt ihm, wie viel behält TikTok?",
         "Creator: 40 €  |  TikTok: 40 €  (die Hälfte von 80 €)."),
        ("b) Wie viel müssen Fans insgesamt verschenken, damit beim Creator 100 € ankommen?",
         "200 € - denn nur die Hälfte kommt beim Creator an."),
        ("c) Knifflig: Erkläre, warum TikTok an einem einzigen Geschenk sogar DOPPELT verdient. "
         "(Tipp: Wie kommen die Fans überhaupt an ihre Coins?)",
         "TikTok verdient zweimal: (1) beim Verkauf der Coins an die Fans und (2) noch einmal an der "
         "~50-%-Provision, wenn der Creator sich auszahlen lässt."),
    ]
    for frage, antwort in profi:
        pf = add_para(doc, frage, size=11, bold=True, color=DARK, space_after=2, space_before=4)
        if loesungen:
            add_para(doc, "Lösung:  " + antwort, size=10, bold=True, color=GREEN, space_after=4)
        else:
            pf.paragraph_format.keep_with_next = True
            add_lines(doc, 2, keep_together=True)

    # ──────────────────────────────────────────────────────────
    # SELBSTCHECK - Selbsteinschaetzung (Hattie-Hebel)
    # ──────────────────────────────────────────────────────────
    add_para(doc, "Selbstcheck - Wie sicher bin ich?", bold=True, size=12, color=TT_CYAN,
             space_before=14, space_after=4)
    selbstcheck = [
        "Ich kann in einem Satz erklären, was TikTok mit meinen Videos machen darf.",
        "Ich kenne mindestens 2 Daten, die TikTok automatisch über mich sammelt.",
        "Ich kann den 50-%-Trick bei LIVE-Geschenken erklären.",
    ]
    for satz in selbstcheck:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(satz + "   ")
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK
        r2 = p.add_run("[ ] sicher   [ ] teils   [ ] noch unsicher")
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Das nehme ich mir ab heute konkret vor:  ")
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = DARK
    r2 = p.add_run("_______________________________________________")
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ──────────────────────────────────────────────────────────
    # FOOTER
    # ──────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Quellen: TikTok-Nutzungsbedingungen (EU) - TikTok-Datenschutzrichtlinie - "
                  "Irish Data Protection Commission (2025) - Verbraucherzentrale.de")
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = GREY

    add_page_number_footer(doc)

    return doc


if __name__ == "__main__":
    import os
    out_dir = os.path.dirname(os.path.abspath(__file__))

    ab = baue_arbeitsblatt(loesungen=False)
    ab_path = f"{out_dir}/Arbeitsblatt_TikTok_AGB.docx"
    ab.save(ab_path)
    print(f"OK Arbeitsblatt gespeichert: {ab_path}")

    lo = baue_arbeitsblatt(loesungen=True)
    lo_path = f"{out_dir}/Arbeitsblatt_TikTok_AGB_Loesungen.docx"
    lo.save(lo_path)
    print(f"OK Loesungsblatt gespeichert: {lo_path}")
